import hashlib
import io
import re
from typing import Dict, List


def load_document(file) -> str:
    """Load PDF, DOCX, or TXT content from a Streamlit-uploaded file."""
    filename = file.name.lower()
    raw = file.read()

    if filename.endswith(".pdf"):
        return _load_pdf(raw)

    if filename.endswith(".docx"):
        from docx import Document

        document = Document(io.BytesIO(raw))
        return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()

    if filename.endswith(".txt"):
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return raw.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore").strip()

    raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")


def _load_pdf(raw: bytes) -> str:
    """Extract text from PDF bytes with graceful fallbacks."""
    extractors = (_load_pdf_with_pymupdf, _load_pdf_with_pypdfium2, _load_pdf_with_pdfplumber)
    errors = []

    for extractor in extractors:
        try:
            text = extractor(raw)
            if text.strip():
                return text.strip()
        except ImportError as exc:
            errors.append(str(exc))
        except Exception as exc:
            errors.append(f"{extractor.__name__}: {str(exc)}")

    raise ValueError(
        "Could not extract text from this PDF. If it is scanned, please OCR it first. "
        f"Extractor details: {'; '.join(errors)}"
    )


def _load_pdf_with_pymupdf(raw: bytes) -> str:
    import fitz

    text_parts = []
    with fitz.open(stream=raw, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def _load_pdf_with_pypdfium2(raw: bytes) -> str:
    import pypdfium2 as pdfium

    text_parts = []
    pdf = pdfium.PdfDocument(raw)
    try:
        for page in pdf:
            textpage = page.get_textpage()
            text_parts.append(textpage.get_text_bounded())
            textpage.close()
            page.close()
    finally:
        pdf.close()
    return "\n".join(text_parts)


def _load_pdf_with_pdfplumber(raw: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts)


def _sentences(text: str) -> List[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []
    return re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", normalized)


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Sentence-aware chunking with a small token overlap."""
    sentences = _sentences(text)
    chunks: List[str] = []
    current: List[str] = []
    current_words = 0

    for sentence in sentences:
        sentence_words = sentence.split()
        if current and current_words + len(sentence_words) > chunk_size:
            chunks.append(" ".join(current).strip())
            overlap_words = " ".join(current).split()[-overlap:] if overlap else []
            current = [" ".join(overlap_words)] if overlap_words else []
            current_words = len(overlap_words)

        if len(sentence_words) > chunk_size:
            words = sentence_words
            step = max(1, chunk_size - overlap)
            for start in range(0, len(words), step):
                window = words[start : start + chunk_size]
                chunks.append(" ".join(window).strip())
            current = []
            current_words = 0
        else:
            current.append(sentence)
            current_words += len(sentence_words)

    if current:
        chunks.append(" ".join(current).strip())

    return [chunk for chunk in chunks if chunk]


def embed_and_store(chunks: List[str], collection_name: str, metadata: Dict) -> int:
    """Embed chunks and store them in Chroma using the shared manager resources."""
    if not chunks:
        return 0

    try:
        from .bm25_index import clear_bm25_cache
        from .multi_collection_manager import multi_collection_manager
    except ImportError:
        from bm25_index import clear_bm25_cache
        from multi_collection_manager import multi_collection_manager

    collection = multi_collection_manager.client.get_or_create_collection(collection_name)
    embeddings = multi_collection_manager.model.encode(chunks).tolist()
    document_name = metadata.get("document_name", "uploaded_document")
    digest = hashlib.sha1(f"{document_name}:{len(chunks)}:{chunks[0][:80]}".encode("utf-8")).hexdigest()[:12]

    ids = [f"{collection_name}_{digest}_{index}" for index in range(len(chunks))]
    metadatas = []
    for index, _ in enumerate(chunks):
        chunk_meta = {
            "act": "Uploaded Document",
            "document_name": document_name,
            "source": "user_upload",
            "chunk_index": index,
            "section": f"upload-{index + 1}",
            "section_display": f"Uploaded chunk {index + 1}",
            "heading": metadata.get("heading", document_name),
        }
        metadatas.append(chunk_meta)

    collection.upsert(
        ids=ids,
        documents=chunks,
        embeddings=embeddings,
        metadatas=metadatas,
    )
    clear_bm25_cache()
    return len(chunks)
